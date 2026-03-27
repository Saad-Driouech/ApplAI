"""
Cover letter generator.

Calls the Anthropic API to produce a tailored cover letter,
then writes it as a .docx file using python-docx.

The candidate profile is read from profile_summary.md on every call
so updates to the file are picked up automatically.

Output: {job_folder}/cover_letter.docx
"""
from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Optional

from src.logger import audit, get_logger
from src.utils.sanitize import sanitize_company_for_path

log = get_logger(__name__)

_CL_SYSTEM_TEMPLATE = """\
You are an expert cover letter writer for the candidate described below.
Never fabricate qualifications or experiences not present in the profile.

## Candidate Profile
{profile}

## Cover Letter Rules
- Maximum 400 words.
- Three paragraphs: hook/connection to the company, evidence of fit, call-to-action.
- Tone: professional, confident, concise.
- Use specific metrics and achievements from the profile where relevant.
- Tailor to the company's specific mission and focus area.
- Do NOT include a salutation line (added separately by the formatter).
- Return ONLY the letter body text — no markdown, no subject line, no sign-off.
"""

_CL_USER_TEMPLATE = """\
Write a cover letter for this role:

Title: {title}
Company: {company}
Country: {country}
Scoring Analysis:
{reasoning}

Job Description (first 2000 chars):
{description}
"""


class CoverLetterError(Exception):
    pass


class CoverLetterGenerator:
    """
    Generates a tailored .docx cover letter for a specific job.

    Args:
        output_dir:      Base directory where job folders are created.
        profile_path:    Path to profile_summary.md — read fresh on every call.
        candidate_name:  Used in the sign-off of the letter.
        candidate_email: Included in the document header.
    """

    def __init__(
        self,
        output_dir: Path,
        profile_path: Optional[Path],
        candidate_name: str,
        candidate_email: str,
    ):
        self._output_dir = output_dir
        self._profile_path = profile_path
        self._candidate_name = candidate_name
        self._candidate_email = candidate_email

    def generate(self, job: dict, gemini_reasoning: str = "") -> Path:
        """
        Generate a .docx cover letter for *job*.

        Returns the path to the .docx file.
        Raises CoverLetterError on failure.
        """
        body = self._write_with_api(job, gemini_reasoning)
        job_folder = self._job_folder(job)
        job_folder.mkdir(parents=True, exist_ok=True)

        docx_path = self._save_docx(body, job, job_folder)
        audit("cover_letter_generated", job_id=job.get("id", "?"), docx_path=str(docx_path))
        log.info("Cover letter generated: %s", docx_path)
        return docx_path

    def _write_with_api(self, job: dict, reasoning: str) -> str:
        from src import claude_bridge

        data = claude_bridge.validate_input(dict(job))

        # Read fresh on every call — profile_summary.md changes over time
        profile = (
            self._profile_path.read_text(encoding="utf-8")
            if self._profile_path and self._profile_path.exists()
            else ""
        )
        if not profile:
            log.warning(
                "profile_summary.md not found at %s — cover letter will lack candidate context",
                self._profile_path,
            )

        system = _CL_SYSTEM_TEMPLATE.format(profile=profile)
        user = _CL_USER_TEMPLATE.format(
            title=data["title"],
            company=data["company"],
            country=data["country"],
            reasoning=reasoning[:1000],
            description=(data.get("description") or "")[:2000],
        )

        result = claude_bridge.call_api(system, user, max_tokens=1024)

        if result.get("status") != "success":
            raise CoverLetterError(f"API error: {result.get('error', 'unknown')}")

        output = result.get("output", "").strip()
        if not output:
            raise CoverLetterError("API returned empty cover letter")

        return output

    def _save_docx(self, body: str, job: dict, job_folder: Path) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise CoverLetterError(
                "python-docx is not installed. Run: pip install python-docx"
            ) from exc

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run(self._candidate_name)
        run.bold = True
        run.font.size = Pt(12)
        header.add_run(f"\n{self._candidate_email}")

        doc.add_paragraph()

        salutation = doc.add_paragraph()
        salutation.add_run(f"Hiring Manager\n{job.get('company', '')}")

        doc.add_paragraph()

        re_line = doc.add_paragraph()
        re_run = re_line.add_run(f"Re: {job.get('title', 'Application')}")
        re_run.bold = True

        doc.add_paragraph()

        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(8)

        doc.add_paragraph()

        signoff = doc.add_paragraph()
        signoff.add_run("Sincerely,\n")
        signoff_name = signoff.add_run(self._candidate_name)
        signoff_name.bold = True

        docx_path = job_folder / "cover_letter.docx"
        doc.save(str(docx_path))
        return docx_path

    def _job_folder(self, job: dict) -> Path:
        company = sanitize_company_for_path(job.get("company", "unknown"))
        title_slug = re.sub(r"[^\w\-]", "_", job.get("title", "job"))[:40]
        uid = job.get("id", str(uuid.uuid4()))[:8]
        return self._output_dir / f"{company}_{title_slug}_{uid}"
